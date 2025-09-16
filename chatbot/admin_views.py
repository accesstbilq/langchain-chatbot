# =============================
# Django & Library Imports
# =============================
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import os
from dotenv import load_dotenv
from datetime import datetime
import mimetypes
import uuid
from .models import ChatSession, ChatMessage, Document
from django.http import FileResponse, Http404
import chromadb
from chromadb.config import Settings
import openai
from typing import List, Dict, Any
import json
import PyPDF2
import docx
import csv
import openpyxl
from pptx import Presentation
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
from django.utils import timezone
import os
import re

from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument


# -------------------------------
# Global variables
# -------------------------------
chat_system = ""     # system prompt storage
messages = []        # in-memory messages list

# -------------------------------
# Environment setup
# -------------------------------
load_dotenv()
OPENAIKEY = os.getenv('OPEN_AI_KEY')
MODEL = "gpt-4.1"
EMBEDDING_MODEL = "text-embedding-3-small"
APP_URL = "http://127.0.0.1:8000/"
# Initialize OpenAI client
openai.api_key = OPENAIKEY

# -------------------------------
# Chroma DB setup
# -------------------------------
CHROMA_DB_PATH = os.path.join(settings.BASE_DIR, "chroma_db")
if not os.path.exists(CHROMA_DB_PATH):
    os.makedirs(CHROMA_DB_PATH)

embedding_function = OpenAIEmbeddings(model=EMBEDDING_MODEL, openai_api_key=OPENAIKEY)

vectorstore = Chroma(
    collection_name="documents",
    embedding_function=embedding_function,
    persist_directory=CHROMA_DB_PATH,
)


# ==============================
# 2. Split into Chunks
# ==============================
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1500,
    chunk_overlap=200,
    separators=["\n\n", "\n"]
)

# -------------------------------
# Document storage configuration
# -------------------------------
UPLOAD_DIR = os.path.join(settings.MEDIA_ROOT, 'documents')
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)


def extract_text_and_split(file_path, file_extension, base_metadata):
    """
    Extract text based on file type and split by paragraphs.
    """
    docs = []

    # PDF
    if file_extension == "pdf":
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        for i, page in enumerate(pages):
            paragraphs = [p.strip() for p in page.page_content.split("\n\n") if p.strip()]
            for j, para in enumerate(paragraphs):
                meta = base_metadata.copy()
                meta.update({
                    "page": i,
                    "chunk_index": j,
                    "chunk_type": "paragraph",
                    "chunk_size": len(para),
                })
                docs.append(LCDocument(page_content=para, metadata=meta))

    # DOCX
    elif file_extension == "docx":
        loader = Docx2txtLoader(file_path)
        pages = loader.load()
        text = " ".join([p.page_content for p in pages])
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for i, para in enumerate(paragraphs):
            meta = base_metadata.copy()
            meta.update({
                "chunk_index": i,
                "chunk_type": "paragraph",
                "chunk_size": len(para),
            })
            docs.append(LCDocument(page_content=para, metadata=meta))

    # TXT
    elif file_extension == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        for i, para in enumerate(paragraphs):
            meta = base_metadata.copy()
            meta.update({
                "chunk_index": i,
                "chunk_type": "paragraph",
                "chunk_size": len(para),
            })
            docs.append(LCDocument(page_content=para, metadata=meta))

    # CSV (each row as a paragraph)
    elif file_extension == "csv":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            rows = f.readlines()
        for i, row in enumerate(rows):
            meta = base_metadata.copy()
            meta.update({
                "chunk_index": i,
                "chunk_type": "row",
                "chunk_size": len(row),
            })
            docs.append(LCDocument(page_content=row.strip(), metadata=meta))

    return docs


def save_docs_to_json_and_chroma(docs, base_dir="documents", filename=None):
    """
    Save docs into JSON file and also into ChromaDB with extra metadata.
    """
    # Ensure folder exists
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)

    # Generate unique JSON filename if not provided
    if not filename:
        json_id = str(uuid.uuid4())
        filename = f"chunk_data_{json_id}.json"

    output_file = os.path.join(base_dir, filename)

    enriched_docs = []
    data = []
    for doc in docs:
        meta = doc.metadata.copy()
        meta["json_file"] = filename   # ✅ add JSON file name to metadata
        data.append({
            "metadata": meta,
            "page_content": doc.page_content
        })
        enriched_docs.append(LCDocument(page_content=doc.page_content, metadata=meta))

    # Save JSON file
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

    print(f"✅ Saved {len(docs)} chunks into {output_file}")

    # Save into Chroma
    vectorstore.add_documents(enriched_docs)
    print(f"✅ Stored {len(enriched_docs)} chunks in ChromaDB")

    return output_file



# -------------------------------
# Chat history view
# -------------------------------
def chathistory(request):
    """Render chat history page with all sessions and last session messages."""
    sessions, last_session_id = load_chat_session() or ([], None)

    # if no last session, set messages empty
    if last_session_id:
        messages_qs = load_chat_history_from_db(last_session_id) or []
    else:
        messages_qs = []

    return render(request, 'chathistory.html', {
        'sessions': sessions,
        'last_session_id': last_session_id,
        'messages': messages_qs,
    })

# -------------------------------
# Document view
# -------------------------------
def document_view(request):
    """Render document upload page."""
    return render(request, 'document.html')

# -------------------------------
# Upload documents with Chroma integration
# -------------------------------
@csrf_exempt
def upload_documents(request):
    """Handle multiple document uploads with validation, metadata storage, and embedding generation."""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Only POST requests allowed'})
    
    if 'documents' not in request.FILES:
        return JsonResponse({'success': False, 'error': 'No documents provided'})
    
    uploaded_files = []
    errors = []
    
    # Allowed file extensions
    ALLOWED_EXTENSIONS = {
        'txt', 'pdf', 'doc', 'docx', 'html', 'css', 'js', 
        'json', 'xml', 'csv', 'xlsx', 'ppt', 'pptx'
    }
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 10MB
    
    files = request.FILES.getlist('documents')
    session_id = request.POST.get('session_id')  # Optional session association
    
    # Get session object if provided
    session = None
    if session_id:
        try:
            session = ChatSession.objects.get(session_id=session_id)
        except ChatSession.DoesNotExist:
            pass
    
    for file in files:
        try:
            # Validate file size
            if file.size > MAX_FILE_SIZE:
                errors.append(f"File '{file.name}' is too large (max 50MB)")
                continue
            
            # Validate file extension
            file_extension = file.name.split('.')[-1].lower()
            if file_extension not in ALLOWED_EXTENSIONS:
                errors.append(f"File type '{file_extension}' not allowed for '{file.name}'")
                continue
            
            # Generate unique filename
            unique_id = uuid.uuid4()
            file_name = f"{unique_id}_{file.name}"
            
            # Save file on disk
            fs = FileSystemStorage(location=UPLOAD_DIR)
            filename = fs.save(file_name, file)
            file_path = fs.path(filename)
            
            # loader = PyPDFLoader(file_path)
            # chunk_documents = loader.load()
            

            base_metadata = {
                "document_id": str(unique_id),
                "session_id": str(session.session_id) if session else None,
                "original_name": file.name,
                "file_name": filename,
                "file_size": file.size,
                "file_type": file_extension,
                "mime_type": mimetypes.guess_type(file.name)[0] or "application/octet-stream",
                "file_path": file_path,
                "url": fs.url(filename) if hasattr(fs, "url") else None,
                "upload_date": timezone.now().isoformat()
            }


            docs = extract_text_and_split(file_path, file_extension, base_metadata)
            savefile = save_docs_to_json_and_chroma(docs, base_dir="documents")

            print(f"Split into {len(docs)} chunks with metadata")
            
            # Prepare response entry
            uploaded_files.append({
                'id': unique_id,
                'name': file.name,
                'size': file.size,
                'type': file_extension,
                'savefile': savefile,
                'upload_date': datetime.now().isoformat(),
            })

        except Exception as e:
            errors.append(f"Error uploading '{file.name}': {str(e)}")
    
    return JsonResponse({
        'success': len(uploaded_files) > 0,
        'uploaded_files': uploaded_files,
        'errors': errors,
        'total_uploaded': len(uploaded_files),
        'total_errors': len(errors)
    })

# -------------------------------
# List uploaded documents
# -------------------------------
@csrf_exempt
def list_documents(request):
    """Return list of uploaded documents from Chroma with chapter information."""
    if request.method != 'GET':
        return JsonResponse({'success': False, 'error': 'Only GET requests allowed'})
    
    try:
        # Fetch all documents metadata from Chroma
        all_docs = vectorstore._collection.get(include=["metadatas"])   

        # Dict to keep only one entry per document_id with chapter count
        unique_docs = {}

        for metadata in all_docs.get("metadatas", []):
            if metadata:
                doc_id = metadata.get("document_id")
                if doc_id:
                    if doc_id not in unique_docs:
                        unique_docs[doc_id] = {
                            "id": doc_id,
                            "name": metadata.get("original_name", ""),
                            "size": metadata.get("file_size", 0),
                            "type": metadata.get("file_type", ""),
                            "upload_date": metadata.get("upload_date", ""),
                            "mime_type": metadata.get("mime_type", "application/octet-stream"),
                            "json_file": metadata.get("json_file", ""),
                            "chunk_count": 0,
                            "chunk_types": set(),
                        }
                    
                    # Count chunks and track chunk types
                    unique_docs[doc_id]["chunk_count"] += 1
                    chunk_type = metadata.get("chunk_type", "unknown")
                    unique_docs[doc_id]["chunk_types"].add(chunk_type)

        # Convert set to list for JSON serialization
        for doc in unique_docs.values():
            doc["chunk_types"] = list(doc["chunk_types"])

        # Convert dict → list of unique docs
        results = list(unique_docs.values())

        return JsonResponse({
            'success': True,
            'documents': results,
            'total_count': len(results)
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f"Error fetching documents: {str(e)}",
            'documents': [],
            'total_count': 0
        })

# -------------------------------
# Delete document (updated for Chroma)
# -------------------------------
@csrf_exempt
def delete_document(request):
    """Delete an uploaded document from LangChain Chroma by document_id."""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Only POST requests allowed'})
    
    document_id = request.POST.get('document_id', '').strip()
    if not document_id:
        return JsonResponse({'success': False, 'error': 'Document ID required'})
    
    try:
        original_name = None
        chunk_count = 0

        # Fetch metadata for this document before deleting
        try:
            chunk_results = vectorstore._collection.get(
                where={"document_id": str(document_id)},
                include=["metadatas"]
            )
            chunk_count = len(chunk_results.get("metadatas", []))
            for metadata in chunk_results.get("metadatas", []):
                if metadata and "original_name" in metadata:
                    original_name = metadata["original_name"]
                    break   # first occurrence is enough
        except Exception as e:
            print(f"Error fetching metadata from Chroma: {str(e)}")

        # Delete from Chroma
        try:
            vectorstore._collection.delete(where={"document_id": str(document_id)})
        except Exception as e:
            print(f"Error deleting from Chroma: {str(e)}")
        
        if not original_name:
            original_name = "Unknown"

        return JsonResponse({
            'success': True,
            'message': f"Document '{original_name}' with {chunk_count} chunks deleted successfully from Chroma"
        })
    
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f"Error deleting document: {str(e)}"
        })

# -------------------------------
# Chat session helpers
# -------------------------------
def load_chat_session():
    """Fetch all chat sessions and return with first session id (if exists)."""
    sessions = ChatSession.objects.all().order_by("-created_at")
    if sessions.exists():
        return sessions, sessions.first().session_id
    return [], None   # always return tuple

# -------------------------------
# Chat History helpers
# -------------------------------
def load_chat_history(request):
    """Load messages for a given session id (used for chat UI)."""
    session_id = request.GET.get("session_id")
    messages = load_chat_history_from_db(session_id)
    return render(request, "chat_messages.html", {"messages": messages})

# -------------------------------
# Load Message using session id
# -------------------------------
def load_chat_history_from_db(session_id):
    """Fetch chat history (human + ai messages only) from DB."""
    if not session_id:
        return []
    session = ChatSession.objects.get(session_id=session_id)

    # Filter messages with non-empty content
    messages = ChatMessage.objects.filter(
        session=session,
        message_type__in=['human', 'ai']  # only include human/ai
    ).exclude(content__isnull=True).exclude(content__exact='').order_by('-created_at')
    
    return messages

def download_document(request, doc_id):
    """Serve uploaded documents by ID"""
    try:
        document = Document.objects.get(document_id=doc_id)
        file_path = document.file_path
        
        response = FileResponse(
            open(file_path, "rb"), 
            as_attachment=True, 
            filename=document.original_name
        )
        
        return response
        
    except Document.DoesNotExist:
        raise Http404("Document not found")
    except FileNotFoundError:
        raise Http404("File not found on server")
    
def chunk_document(request, filename):
    file_path = os.path.join(settings.BASE_DIR, "documents", filename)
    if os.path.exists(file_path):
        return FileResponse(open(file_path, "rb"), as_attachment=True, filename=filename)
    raise Http404("File not found")
